#!/usr/bin/env python3
"""Build the venue-neutral Wave 4 academic manuscript as a styled DOCX.

The Markdown manuscript remains the textual source of truth. This builder only
translates its frozen prose, tables, and simple diagrams into an editable Word
document; it does not calculate, score, or reinterpret any study result.
"""

from __future__ import annotations

import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "manuscript.md"
OUTPUT = ROOT / "manuscript.docx"
SKILL_ROOT = Path(
    "/Users/marbin/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.826.12353/skills/documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


# Preset: narrative_proposal. Named academic-manuscript overrides are limited
# to the title block, running furniture, captions, quotations, references, and
# compact table text. Body and heading tokens retain the selected preset.
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_BODY = "Calibri"
FONT_MONO = "Consolas"
COLOR_NAVY = "0B2545"
COLOR_BLUE = "2E74B5"
COLOR_DARK_BLUE = "1F4D78"
COLOR_GRAY = "5B6573"
COLOR_LIGHT_GRAY = "F4F6F9"
COLOR_TABLE_HEADER = "E8EEF5"
COLOR_BORDER = "B9C2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = COLOR_BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, name: str, size: float, color: str | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, FONT_BODY, 9)
    run.font.color.rgb = RGBColor.from_string(COLOR_GRAY)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_BODY)
    r_fonts.set(qn("w:hAnsi"), FONT_BODY)
    run_props.extend([r_fonts, color, underline])
    new_run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|<https?://[^>]+>|https?://[^\s>]+)"
)


def add_inline(paragraph, text: str, *, base_size: float | None = None) -> None:
    """Add the small Markdown subset used by the manuscript."""
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, FONT_BODY, base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, FONT_BODY, base_size)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, FONT_MONO, (base_size or 11) - 0.5)
            run.font.color.rgb = RGBColor.from_string(COLOR_DARK_BLUE)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            set_run_font(run, FONT_BODY, base_size)
        else:
            url = token[1:-1] if token.startswith("<") else token.rstrip(".,;)")
            suffix = token[len(url) + (1 if token.startswith("<") else 0) :]
            add_hyperlink(paragraph, url, url)
            if suffix and not token.startswith("<"):
                run = paragraph.add_run(suffix)
                set_run_font(run, FONT_BODY, base_size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, FONT_BODY, base_size)


def add_border_bottom(paragraph, color: str = COLOR_BORDER, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_BODY, 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COLOR_BLUE, 18, 10),
        ("Heading 2", 13, COLOR_BLUE, 12, 6),
        ("Heading 3", 12, COLOR_DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        set_style_font(style, FONT_BODY, size, color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, FONT_BODY, 11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Article Metadata" not in doc.styles:
        meta = doc.styles.add_style("Article Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = doc.styles["Article Metadata"]
    set_style_font(meta, FONT_BODY, 10.5, COLOR_GRAY)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    meta.paragraph_format.line_spacing = 1.08

    if "Table Caption" not in doc.styles:
        caption = doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Table Caption"]
    set_style_font(caption, FONT_BODY, 10, COLOR_DARK_BLUE)
    caption.font.bold = True
    caption.paragraph_format.space_before = Pt(10)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    if "Code Block" not in doc.styles:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    set_style_font(code, FONT_MONO, 9.5, COLOR_NAVY)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    if "Reference" not in doc.styles:
        reference = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = doc.styles["Reference"]
    set_style_font(reference, FONT_BODY, 10.5)
    reference.paragraph_format.left_indent = Inches(0.3)
    reference.paragraph_format.first_line_indent = Inches(-0.3)
    reference.paragraph_format.space_after = Pt(5)
    reference.paragraph_format.line_spacing = 1.15

    if "Block Quote" not in doc.styles:
        quote = doc.styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = doc.styles["Block Quote"]
    set_style_font(quote, FONT_BODY, 10.5, COLOR_NAVY)
    quote.font.italic = True
    quote.paragraph_format.left_indent = Inches(0.4)
    quote.paragraph_format.right_indent = Inches(0.3)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "Consumer AI advice under unresolved port-access conditions"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(3)
    for run in header.runs:
        set_run_font(run, FONT_BODY, 8.5)
        run.font.color.rgb = RGBColor.from_string(COLOR_GRAY)
    add_border_bottom(header, color="D7DBE2", size="4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    lead = footer.add_run("Page ")
    set_run_font(lead, FONT_BODY, 9)
    lead.font.color.rgb = RGBColor.from_string(COLOR_GRAY)
    add_page_field(footer)

    doc.core_properties.title = (
        "Consumer AI Dispatch Advice Under Unresolved Port-Access Conditions: "
        "A 96-Response Challenge-Set Evaluation"
    )
    doc.core_properties.author = "Authors to be supplied"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.created = datetime(2026, 8, 31, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 8, 31, tzinfo=timezone.utc)
    doc.core_properties.subject = "Working academic manuscript"
    doc.core_properties.keywords = (
        "consumer artificial intelligence; drayage; port access; dispatch; "
        "human evaluation; Cohen's kappa"
    )
    doc.core_properties.comments = (
        "Generated from manuscript.md without changing frozen study content."
    )


def normalize_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def table_widths(headers: list[str]) -> list[int]:
    joined = " | ".join(headers).lower()
    count = len(headers)
    if count == 5 and "scenario focus" in joined:
        return [600, 1500, 3000, 2100, 2160]
    if count == 5 and "scenario" in joined:
        return [3500, 950, 950, 1660, 2300]
    if count == 5:
        return [3000, 950, 950, 1760, 2700]
    if count == 4 and "frozen construct" in joined:
        return [2400, 1600, 1250, 4110]
    if count == 4 and "endpoint-positive share" in joined:
        return [4560, 1400, 1400, 2000]
    if count == 2 and "sha-256" in joined:
        return [3200, 6160]
    if count == 2:
        return [6100, 3260]
    if count == 3:
        return [3000, 3000, 3360]
    base = PAGE_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, raw_rows: list[str]) -> None:
    rows = [normalize_table_row(line) for line in raw_rows]
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return
    col_count = len(rows[0])
    if any(len(row) != col_count for row in rows):
        raise ValueError("Markdown table contains inconsistent column counts")

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = table_widths(rows[0])
    is_hash_table = col_count == 2 and "SHA-256" in rows[0]
    text_size = 8.2 if (col_count == 5 and "Scenario focus" in rows[0]) else 9.0
    if col_count == 2:
        text_size = 9.5

    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if col_index > 0 and len(value) <= 28
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            cell_text_size = 7.5 if is_hash_table and col_index == 1 else text_size
            add_inline(paragraph, value, base_size=cell_text_size)
            if row_index == 0:
                set_cell_shading(cell, COLOR_TABLE_HEADER)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(COLOR_NAVY)
        if row_index == 0:
            set_repeat_table_header(row)

    apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa={"top": 90, "bottom": 90, "start": 120, "end": 120},
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    after.paragraph_format.keep_with_next = False


def shade_paragraph(paragraph, fill: str = COLOR_LIGHT_GRAY) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_code_block(doc: Document, lines: list[str], figure: bool = False) -> None:
    for idx, line in enumerate(lines):
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = idx < len(lines) - 1
        shade_paragraph(paragraph)
        run = paragraph.add_run(line)
        set_run_font(run, FONT_MONO, 9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_title_page(doc: Document, title: str, metadata_lines: list[str]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(34)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("ORIGINAL RESEARCH · WORKING MANUSCRIPT")
    set_run_font(run, FONT_BODY, 10)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_BLUE)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(18)
    title_paragraph.paragraph_format.keep_together = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, FONT_BODY, 24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(COLOR_NAVY)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.left_indent = Inches(1.4)
    rule.paragraph_format.right_indent = Inches(1.4)
    rule.paragraph_format.space_after = Pt(20)
    add_border_bottom(rule, color=COLOR_BLUE, size="7")

    for line in metadata_lines:
        paragraph = doc.add_paragraph(style="Article Metadata")
        add_inline(paragraph, line, base_size=10.5)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(22)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run(
        "Administrative placeholders are intentionally retained for author completion."
    )
    set_run_font(note_run, FONT_BODY, 9)
    note_run.italic = True
    note_run.font.color.rgb = RGBColor.from_string(COLOR_GRAY)

def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Expected a level-one manuscript title")

    title = lines[0][2:].strip()
    cursor = 1
    metadata: list[str] = []
    while cursor < len(lines):
        line = lines[cursor].strip()
        if line.startswith("## "):
            break
        if line:
            metadata.append(line)
        cursor += 1
    add_title_page(doc, title, metadata)

    in_references = False
    last_caption_was_figure = False
    while cursor < len(lines):
        raw = lines[cursor]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            cursor += 1
            continue

        if stripped.startswith("```"):
            cursor += 1
            block: list[str] = []
            while cursor < len(lines) and not lines[cursor].strip().startswith("```"):
                block.append(lines[cursor].rstrip())
                cursor += 1
            add_code_block(doc, block, figure=last_caption_was_figure)
            last_caption_was_figure = False
            cursor += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while cursor < len(lines) and lines[cursor].strip().startswith("|"):
                table_lines.append(lines[cursor].strip())
                cursor += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, stripped[4:])
            cursor += 1
            continue

        if stripped.startswith("## "):
            heading = stripped[3:]
            paragraph = doc.add_paragraph(style="Heading 1")
            if heading in {
                "Abstract",
                "References",
                "Appendix A. Protocol-to-report crosswalk",
            }:
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, heading)
            in_references = heading == "References"
            cursor += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            cursor += 1
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph(style="Block Quote")
            add_inline(paragraph, stripped[2:], base_size=10.5)
            cursor += 1
            continue

        if stripped.startswith("**Table ") or stripped.startswith("**Figure "):
            paragraph = doc.add_paragraph(style="Table Caption")
            add_inline(paragraph, stripped, base_size=10)
            last_caption_was_figure = stripped.startswith("**Figure ")
            cursor += 1
            continue

        if in_references and re.match(r"^\d+\.\s", stripped):
            paragraph = doc.add_paragraph(style="Reference")
            add_inline(paragraph, stripped, base_size=10.5)
            cursor += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        cursor += 1


def audit_document(doc: Document) -> None:
    if len(doc.tables) != 8:
        raise ValueError(f"Expected 8 manuscript tables; found {len(doc.tables)}")
    if not any("15 of 96" in paragraph.text for paragraph in doc.paragraphs):
        raise ValueError("Primary result text was not preserved")
    if not any("Cohen’s κ was 0.824" in paragraph.text for paragraph in doc.paragraphs):
        raise ValueError("Agreement statistic text was not preserved")
    title_count = sum(
        1
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("Consumer AI Dispatch Advice")
    )
    if title_count != 1:
        raise ValueError(f"Expected one title paragraph; found {title_count}")


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    parse_markdown(doc, markdown)
    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
