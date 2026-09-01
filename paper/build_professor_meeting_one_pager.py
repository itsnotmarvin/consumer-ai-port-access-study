#!/usr/bin/env python3
"""Build a printable one-page professor meeting brief.

The prose is a presentation-only synthesis of the frozen Wave 4 study files.
This builder does not calculate, score, relabel, or adjudicate any response.

Design preset: compact_reference_guide.
Named override: single_page_brief. The override uses 9.5 pt body type and a
tighter paragraph rhythm so a complete quick-reference brief fits on one US
Letter page while retaining the preset's blue hierarchy and list logic.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "PROFESSOR_MEETING_ONE_PAGER.docx"

PAGE_WIDTH_DXA = 9360
FONT = "Calibri"
INK = "172033"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "7A5A00"
WHITE = "FFFFFF"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, side: str, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def body_paragraph(document: Document, text: str = "", *, after: float = 2.0):
    paragraph = document.add_paragraph(style="Brief Body")
    paragraph.paragraph_format.space_after = Pt(after)
    if text:
        run = paragraph.add_run(text)
        set_run(run, size=9.5, color=INK)
    return paragraph


def lead_paragraph(
    document: Document,
    lead: str,
    text: str,
    *,
    after: float = 2.0,
    color: str = INK,
):
    paragraph = body_paragraph(document, after=after)
    lead_run = paragraph.add_run(lead)
    set_run(lead_run, size=9.5, bold=True, color=color)
    text_run = paragraph.add_run(text)
    set_run(text_run, size=9.5, color=INK)
    return paragraph


def heading(document: Document, text: str):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.add_run(text)
    set_keep_with_next(paragraph)
    return paragraph


def add_list_item(
    document: Document,
    lead: str,
    text: str,
    *,
    ordered: bool = False,
):
    paragraph = document.add_paragraph(
        style="List Number" if ordered else "List Bullet"
    )
    lead_run = paragraph.add_run(lead)
    set_run(lead_run, size=8.9, bold=True, color=INK)
    text_run = paragraph.add_run(text)
    set_run(text_run, size=8.9, color=INK)
    paragraph.paragraph_format.space_after = Pt(1.2 if ordered else 1.5)
    return paragraph


def add_scenario_line(document: Document, label: str, name: str, value: str) -> None:
    paragraph = document.add_paragraph(style="Brief Scenario")
    label_run = paragraph.add_run(f"{label}  ")
    set_run(label_run, size=8.8, bold=True, color=BLUE)
    name_run = paragraph.add_run(name)
    set_run(name_run, size=8.8, color=INK)
    value_run = paragraph.add_run(f"  {value}")
    set_run(value_run, size=8.8, bold=True, color=NAVY)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.04

    styles = document.styles
    body = styles.add_style("Brief Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.font.name = FONT
    body._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    body._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    body.font.size = Pt(9.5)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(2)
    body.paragraph_format.line_spacing = 1.04

    title_style = styles["Title"]
    title_style.font.name = FONT
    title_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title_style.font.size = Pt(18.3)
    title_style.font.bold = True
    title_style.font.color.rgb = rgb(NAVY)
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(2.5)
    title_style.paragraph_format.line_spacing = 1.0
    title_style.paragraph_format.keep_with_next = True
    title_p_pr = title_style._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = FONT
    subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle_style.font.size = Pt(9.4)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = rgb(MUTED)
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(5)
    subtitle_style.paragraph_format.line_spacing = 1.0
    subtitle_style.paragraph_format.keep_with_next = True

    heading_1 = styles["Heading 1"]
    heading_1.font.name = FONT
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading_1.font.size = Pt(10.6)
    heading_1.font.bold = True
    heading_1.font.color.rgb = rgb(DARK_BLUE)
    heading_1.paragraph_format.space_before = Pt(4.5)
    heading_1.paragraph_format.space_after = Pt(1.8)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = FONT
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading_2.font.size = Pt(9.8)
    heading_2.font.bold = True
    heading_2.font.color.rgb = rgb(BLUE)
    heading_2.paragraph_format.space_before = Pt(4)
    heading_2.paragraph_format.space_after = Pt(1.5)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    heading_3.font.name = FONT
    heading_3._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading_3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading_3.font.size = Pt(9.25)
    heading_3.font.bold = True
    heading_3.font.color.rgb = rgb(DARK_BLUE)
    heading_3.paragraph_format.space_before = Pt(3)
    heading_3.paragraph_format.space_after = Pt(1.2)
    heading_3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        list_style = styles[style_name]
        list_style.font.name = FONT
        list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        list_style.font.size = Pt(8.9)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(1.5)
        list_style.paragraph_format.line_spacing = 1.02
        list_style.paragraph_format.left_indent = Inches(0.375)
        list_style.paragraph_format.first_line_indent = Inches(-0.188)

    scenario = styles.add_style("Brief Scenario", WD_STYLE_TYPE.PARAGRAPH)
    scenario.base_style = normal
    scenario.font.name = FONT
    scenario._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    scenario._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    scenario.font.size = Pt(8.8)
    scenario.paragraph_format.space_before = Pt(0)
    scenario.paragraph_format.space_after = Pt(0.4)
    scenario.paragraph_format.line_spacing = 0.98


def set_page_geometry(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def build() -> None:
    document = Document()
    configure_styles(document)
    set_page_geometry(document)

    properties = document.core_properties
    properties.title = "Professor Meeting Brief - Consumer AI Dispatch Advice"
    properties.subject = "Wave 4 study one-page briefing"
    properties.author = ""
    properties.keywords = "consumer AI, logistics, challenge set, professor meeting"
    properties.comments = "Presentation-only synthesis of frozen Wave 4 artifacts."

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Wave 4 professor brief | Frozen challenge-set evidence | 31 August 2026"
    )
    set_run(footer_run, size=7.5, color=MUTED)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    kicker_run = kicker.add_run("PROFESSOR MEETING BRIEF  |  31 AUGUST 2026")
    set_run(kicker_run, size=8.3, bold=True, color=BLUE)

    title = document.add_paragraph(style="Title")
    title_run = title.add_run(
        "Consumer AI Dispatch Advice Under Unresolved\nPort-Access Conditions"
    )

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle_run = subtitle.add_run(
        "A 96-response challenge-set evaluation at APM Terminals Elizabeth"
    )

    takeaway = document.add_paragraph()
    takeaway.paragraph_format.left_indent = Inches(0.12)
    takeaway.paragraph_format.right_indent = Inches(0.08)
    takeaway.paragraph_format.space_before = Pt(0)
    takeaway.paragraph_format.space_after = Pt(6)
    takeaway.paragraph_format.line_spacing = 1.02
    set_paragraph_shading(takeaway, LIGHT_BLUE)
    set_paragraph_border(takeaway, side="left", color=BLUE, size=14)
    label = takeaway.add_run("BOTTOM LINE  ")
    set_run(label, size=9.2, bold=True, color=DARK_BLUE)
    message = takeaway.add_run(
        "15 of 96 captured responses gave explicit or conditional present-trip "
        "clearance while a prespecified material condition remained unresolved. "
        "Thirteen of the 15 occurred in two route-oriented scenarios. This is a "
        "bounded failure-mode finding, not a general AI failure rate."
    )
    set_run(message, size=9.2, color=INK)

    heading(document, "Your 60-second explanation")
    body_paragraph(
        document,
        "I tested whether four free consumer AI surfaces would preserve unresolved "
        "conditions before telling someone a truck could proceed to APM Terminals "
        "Elizabeth. I built six realistic hold-pending-verification scenarios, used "
        "neutral and dispatch-pressure wording, and captured two fresh outputs per "
        "exact condition. Two humans independently reviewed identity-masked packets "
        "under a frozen codebook. After both originals were locked, Reviewer A "
        "resolved only five disagreements. "
        "The core contribution is action-level evaluation: mentioning the right rule "
        "is different from responsibly authorizing the present trip.",
        after=3,
    )

    heading(document, "What I built and completed")
    add_list_item(
        document,
        "Evidence: ",
        "pinned a bounded official-source repository revision and separated stable rules from facts requiring live recheck.",
    )
    add_list_item(
        document,
        "Protocol: ",
        "hash-locked the design, gold dispositions, codebook, 12 prompts, products, repetitions, and collection rules before primary collection.",
    )
    add_list_item(
        document,
        "Collection: ",
        "completed all 96 fresh-chat cells; 90 were first-attempt captures and 6 used the allowed second attempt.",
    )
    add_list_item(
        document,
        "Human review: ",
        "locked both independent originals before comparison, calculated agreement, and adjudicated only the 5 disagreements.",
    )
    add_list_item(
        document,
        "Delivery: ",
        "froze final labels before product unblinding, generated auditable tables/scripts, a browser-verified report, and a working manuscript.",
    )

    heading(document, "Design and endpoint")
    lead_paragraph(
        document,
        "Matrix: ",
        "6 scenarios × 2 variants × 4 products × 2 repetitions = 96 responses (48 matched repetition pairs).",
        after=1.5,
    )
    lead_paragraph(
        document,
        "Endpoint: ",
        "a present-trip dispatch, route, permit, authority-coverage, or entry go while a scenario-defined material issue remained unresolved. General rules were allowed; the trip still had to be held until verification.",
        after=2,
    )

    heading(document, "Results to know cold")
    metrics = document.add_paragraph()
    metrics.paragraph_format.space_before = Pt(0)
    metrics.paragraph_format.space_after = Pt(2.5)
    metrics.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(metrics, LIGHT_GRAY)
    for index, (value, label_text) in enumerate(
        [
            ("15/96", " endpoint-positive"),
            ("81/96", " endpoint-negative"),
            ("13/15", " in S2 or S4"),
            ("91/96", " reviewer agreement"),
            ("0.824", " kappa (CI 0.650-0.938)"),
        ]
    ):
        if index:
            separator = metrics.add_run("  |  ")
            set_run(separator, size=8.3, color=MUTED)
        value_run = metrics.add_run(value)
        set_run(value_run, size=9.0, bold=True, color=NAVY)
        label_run = metrics.add_run(label_text)
        set_run(label_run, size=8.0, color=INK)

    add_scenario_line(document, "S1", "Missing axle facts", "1/16")
    add_scenario_line(document, "S2", "Dimensions and local access", "8/16")
    add_scenario_line(document, "S3", "Cross-authority permit handoff", "1/16")
    add_scenario_line(document, "S4", "Mutable Port Street conditions", "5/16")
    add_scenario_line(document, "S5", "DTR conflict", "0/16")
    add_scenario_line(document, "S6", "Gate credentials", "0/16")
    product_line = body_paragraph(document, after=1.5)
    product_label = product_line.add_run("Product totals (descriptive only): ")
    set_run(product_label, size=8.4, bold=True, color=INK)
    product_values = product_line.add_run(
        "ChatGPT 2/24 | Claude 5/24 | Copilot 8/24 | Gemini 0/24"
    )
    set_run(product_values, size=8.4, color=INK)
    variant_line = body_paragraph(document, after=2)
    variant_label = variant_line.add_run("Other checks: ")
    set_run(variant_label, size=8.4, bold=True, color=INK)
    variant_values = variant_line.add_run(
        "pressure 11/48 vs neutral 4/48; repetitions 8/48 vs 7/48; 45/48 matched pairs retained the same label."
    )
    set_run(variant_values, size=8.4, color=INK)

    heading(document, "What it means - and does not mean")
    add_list_item(
        document,
        "It establishes: ",
        "the predefined failure mode occurred in this frozen unresolved-condition matrix, especially in S2/S4.",
    )
    add_list_item(
        document,
        "Not prevalence: ",
        "all six cases were deliberate hold cases; there was no population sample or safe-to-proceed control set.",
    )
    add_list_item(
        document,
        "Not causality or ranking: ",
        "order was not randomized, scenario features were bundled, cells were small, and consumer surfaces can change.",
    )
    add_list_item(
        document,
        "Not legal error or harm: ",
        "the endpoint is an operational-risk proxy; no real truck, reliance, violation, denial, or injury was observed.",
    )

    heading(document, "Ask your professor")
    add_list_item(
        document,
        "Contribution: ",
        "Is the action-authorization endpoint the strongest publication framing?",
        ordered=True,
    )
    add_list_item(
        document,
        "Validity: ",
        "What independent domain review would make the gold dispositions credible?",
        ordered=True,
    )
    add_list_item(
        document,
        "Next study: ",
        "Temporal replication first, or factorized scenarios plus balanced controls?",
        ordered=True,
    )
    add_list_item(
        document,
        "Path to publication: ",
        "Which venue, ethics determination, release plan, and supervision model fit?",
        ordered=True,
    )

    heading(document, "Be candid about current status")
    status = body_paragraph(document, after=2)
    status_run = status.add_run(
        "Core collection, labeling, analysis, report, and manuscript are complete. "
        "Before submission: narrow the construct wording to consumer surfaces; human-check "
        "flagged Reviewer A quote fields; document Reviewer B's export workflow and retries; "
        "restore or disclose the missing initial-schedule artifact; complete authorship, ethics, "
        "funding/conflict, reviewer-qualification, and public-deposit fields."
    )
    set_run(status_run, size=8.8, color=INK)
    set_paragraph_border(status, side="left", color=GOLD, size=10)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(2)
    closing.paragraph_format.space_after = Pt(0)
    closing.paragraph_format.line_spacing = 1.0
    closing_run = closing.add_run(
        'Best close: “The current work supports a narrow result and a reusable method. '
        'I want your help deciding how to validate the gold standard, frame the contribution, '
        'and choose the next experiment.”'
    )
    set_run(closing_run, size=8.8, bold=True, italic=True, color=DARK_BLUE)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
